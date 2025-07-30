import requests
import json
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import locale
import os
from typing import Dict, Any, Optional

class ShopifyInvoiceGenerator:
    def __init__(self, config: Dict[str, str]):
        self.shopify_config = {
            'shop_domain': config['shop_domain'],
            'access_token': config['shopify_access_token'],
            'api_version': '2023-10'
        }

        self.office_config = {
            'client_id': config.get('office_client_id'),
            'client_secret': config.get('office_client_secret'),
            'tenant_id': config.get('office_tenant_id'),
            'access_token': None
        }

        self.company_info = {
            'name': config.get('company_name', 'Tu Empresa'),
            'address': config.get('company_address', 'Dirección de tu empresa'),
            'phone': config.get('company_phone', 'Teléfono'),
            'email': config.get('company_email', 'email@empresa.com'),
            'tax_id': config.get('company_tax_id', 'NIT/RUT')
        }

        # Configurar formato de moneda colombiana
        try:
            locale.setlocale(locale.LC_ALL, 'es_CO.UTF-8')
        except:
            try:
                locale.setlocale(locale.LC_ALL, 'es_ES.UTF-8')
            except:
                pass  # Usar formato por defecto

    def get_office_access_token(self) -> str:
        """Obtener token de acceso para Microsoft Graph API"""
        try:
            token_url = f"https://login.microsoftonline.com/{self.office_config['tenant_id']}/oauth2/v2.0/token"

            data = {
                'client_id': self.office_config['client_id'],
                'client_secret': self.office_config['client_secret'],
                'scope': 'https://graph.microsoft.com/.default',
                'grant_type': 'client_credentials'
            }

            response = requests.post(token_url, data=data)
            response.raise_for_status()

            token_data = response.json()
            self.office_config['access_token'] = token_data['access_token']
            return token_data['access_token']

        except requests.exceptions.RequestException as e:
            print(f"Error obteniendo token de Office: {e}")
            raise

    def get_shopify_order(self, order_id: str) -> Dict[str, Any]:
        """Obtener datos de la orden desde Shopify"""
        try:
            url = f"https://{self.shopify_config['shop_domain']}/admin/api/{self.shopify_config['api_version']}/orders/{order_id}.json"

            headers = {
                'X-Shopify-Access-Token': self.shopify_config['access_token']
            }

            response = requests.get(url, headers=headers)
            response.raise_for_status()

            return response.json()['order']

        except requests.exceptions.RequestException as e:
            print(f"Error obteniendo orden de Shopify: {e}")
            raise

    def format_currency(self, amount: float) -> str:
        """Formatear moneda en pesos colombianos"""
        try:
            return f"${amount:,.0f} COP"
        except:
            return f"${amount:.2f}"

    def format_date(self, date_string: str) -> str:
        """Formatear fecha"""
        try:
            date_obj = datetime.fromisoformat(date_string.replace('Z', '+00:00'))
            return date_obj.strftime('%d/%m/%Y')
        except:
            return date_string

    def create_word_document(self, order_data: Dict[str, Any], filename: str) -> str:
        """Crear documento de Word con la factura"""
        try:
            # Crear nuevo documento
            doc = Document()

            # Configurar márgenes
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # ENCABEZADO - Información de la empresa y factura
            header_table = doc.add_table(rows=1, cols=2)
            header_table.autofit = False
            header_table.columns[0].width = Inches(3.5)
            header_table.columns[1].width = Inches(2.5)

            # Información de la empresa
            company_cell = header_table.cell(0, 0)
            company_para = company_cell.paragraphs[0]
            company_run = company_para.add_run(self.company_info['name'])
            company_run.font.size = Pt(18)
            company_run.font.bold = True

            # Agregar información adicional de la empresa
            company_cell.add_paragraph(self.company_info['address'])
            company_cell.add_paragraph(f"Tel: {self.company_info['phone']}")
            company_cell.add_paragraph(f"Email: {self.company_info['email']}")
            company_cell.add_paragraph(self.company_info['tax_id'])

            # Información de la factura
            invoice_cell = header_table.cell(0, 1)
            invoice_para = invoice_cell.paragraphs[0]
            invoice_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            invoice_run = invoice_para.add_run('FACTURA')
            invoice_run.font.size = Pt(20)
            invoice_run.font.bold = True

            # Agregar información de la factura
            invoice_info_para = invoice_cell.add_paragraph()
            invoice_info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            invoice_info_text = f"No. Orden: {order_data['order_number']}\n"
            invoice_info_text += f"Fecha: {self.format_date(order_data['created_at'])}\n"
            due_date = (datetime.now() + timedelta(days=30)).strftime('%d/%m/%Y')
            invoice_info_text += f"Vencimiento: {due_date}"
            invoice_info_para.add_run(invoice_info_text)

            # Línea separadora
            doc.add_paragraph()

            # INFORMACIÓN DEL CLIENTE
            customer_table = doc.add_table(rows=1, cols=2)
            customer_table.autofit = False
            customer_table.columns[0].width = Inches(3)
            customer_table.columns[1].width = Inches(3)

            # Facturar a
            bill_to_cell = customer_table.cell(0, 0)
            bill_to_title = bill_to_cell.paragraphs[0]
            bill_to_run = bill_to_title.add_run('FACTURAR A:')
            bill_to_run.font.bold = True

            customer = order_data.get('customer', {})
            if customer:
                bill_to_cell.add_paragraph(f"{customer.get('first_name', '')} {customer.get('last_name', '')}")
                bill_to_cell.add_paragraph(customer.get('email', ''))
                if customer.get('phone'):
                    bill_to_cell.add_paragraph(customer['phone'])

            # Enviar a
            ship_to_cell = customer_table.cell(0, 1)
            ship_to_title = ship_to_cell.paragraphs[0]
            ship_to_run = ship_to_title.add_run('ENVIAR A:')
            ship_to_run.font.bold = True

            shipping_address = order_data.get('shipping_address')
            if shipping_address:
                ship_to_cell.add_paragraph(f"{shipping_address.get('first_name', '')} {shipping_address.get('last_name', '')}")
                ship_to_cell.add_paragraph(shipping_address.get('address1', ''))
                if shipping_address.get('address2'):
                    ship_to_cell.add_paragraph(shipping_address['address2'])
                ship_to_cell.add_paragraph(f"{shipping_address.get('city', '')}, {shipping_address.get('province', '')}")
                ship_to_cell.add_paragraph(f"{shipping_address.get('zip', '')} - {shipping_address.get('country', '')}")
            else:
                ship_to_cell.add_paragraph('Información de envío no disponible')

            doc.add_paragraph()

            # TABLA DE PRODUCTOS
            items_table = doc.add_table(rows=1, cols=5)
            items_table.style = 'Table Grid'

            # Encabezados
            headers = ['Producto', 'SKU', 'Cantidad', 'Precio Unit.', 'Total']
            header_cells = items_table.rows[0].cells

            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Hacer encabezados en negrita
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Agregar productos
            for item in order_data['line_items']:
                row_cells = items_table.add_row().cells
                product_name = item['title']
                if item.get('variant_title'):
                    product_name += f"\n{item['variant_title']}"
                row_cells[0].text = product_name

                row_cells[1].text = item.get('sku', 'N/A')
                row_cells[2].text = str(item['quantity'])
                row_cells[3].text = self.format_currency(float(item['price']))
                row_cells[4].text = self.format_currency(float(item['price']) * item['quantity'])

            doc.add_paragraph()

            # TOTALES
            totals_table = doc.add_table(rows=1, cols=2)
            totals_table.autofit = False
            totals_table.columns[0].width = Inches(4)
            totals_table.columns[1].width = Inches(2)

            # Calcular totales
            subtotal = float(order_data['subtotal_price'])
            taxes = float(order_data.get('total_tax', 0))
            shipping = float(order_data.get('total_shipping_price_set', {}).get('shop_money', {}).get('amount', 0))
            total = float(order_data['total_price'])

            totals_cell = totals_table.cell(0, 1)
            totals_para = totals_cell.paragraphs[0]
            totals_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Agregar subtotal
            subtotal_para = totals_cell.add_paragraph(f"Subtotal: {self.format_currency(subtotal)}")
            subtotal_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Agregar envío si existe
            if shipping > 0:
                shipping_para = totals_cell.add_paragraph(f"Envío: {self.format_currency(shipping)}")
                shipping_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Agregar impuestos si existen
            if taxes > 0:
                taxes_para = totals_cell.add_paragraph(f"Impuestos: {self.format_currency(taxes)}")
                taxes_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Agregar total final en negrita
            total_para = totals_cell.add_paragraph()
            total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total_run = total_para.add_run(f"TOTAL: {self.format_currency(total)}")
            total_run.font.bold = True
            total_run.font.size = Pt(14)

            # PIE DE PÁGINA
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.add_run(f"Gracias por su compra | {self.company_info['name']}")

            footer_para2 = doc.add_paragraph('Esta es una factura generada automáticamente')
            footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Guardar documento
            filepath = f"{filename}.docx"
            doc.save(filepath)

            return filepath

        except Exception as e:
            print(f"Error creando documento de Word: {e}")
            raise

    def upload_to_onedrive(self, file_path: str, filename: str) -> Dict[str, str]:
        """Subir documento a OneDrive usando Microsoft Graph API"""
        try:
            if not self.office_config['access_token']:
                self.get_office_access_token()

            # Leer el archivo
            with open(file_path, 'rb') as file:
                file_content = file.read()

            # Subir a OneDrive
            upload_url = f"https://graph.microsoft.com/v1.0/me/drive/root:/{filename}.docx:/content"

            headers = {
                'Authorization': f"Bearer {self.office_config['access_token']}",
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }

            response = requests.put(upload_url, headers=headers, data=file_content)
            response.raise_for_status()

            file_info = response.json()

            return {
                'document_id': file_info['id'],
                'web_url': file_info['webUrl'],
                'download_url': file_info['@microsoft.graph.downloadUrl']
            }

        except requests.exceptions.RequestException as e:
            print(f"Error subiendo a OneDrive: {e}")
            raise

    def generate_invoice(self, order_id: str, save_local: bool = True, upload_to_cloud: bool = False) -> Dict[str, Any]:
        """Función principal para generar la factura"""
        try:
            print('Obteniendo datos de la orden...')
            order_data = self.get_shopify_order(order_id)

            print('Creando documento de Word...')
            filename = f"Factura_{order_data['order_number']}_{int(datetime.now().timestamp())}"
            local_file_path = self.create_word_document(order_data, filename)

            result = {
                'success': True,
                'order_number': order_data['order_number'],
                'local_file_path': local_file_path if save_local else None
            }

            if upload_to_cloud and all(self.office_config.values()):
                print('Subiendo a OneDrive...')
                cloud_info = self.upload_to_onedrive(local_file_path, filename)
                result['cloud_info'] = cloud_info

                # Eliminar archivo local si se subió a la nube y no se quiere guardar local
                if not save_local:
                    os.remove(local_file_path)
                    result['local_file_path'] = None

            print('¡Factura generada exitosamente!')
            return result

        except Exception as e:
            print(f"Error generando factura: {e}")
            return {
                'success': False,
                'error': str(e)
            }